
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


PACKAGE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = PACKAGE_DIR / "MANUSCRIPT_DRAFT.md"
OUTPUT_PATH = (
    PACKAGE_DIR.parents[1]
    / "outputs"
    / "019f65a9-f2cf-76a1-a693-7952915a0b17"
    / "APS_DPV_support_aware_distributional_effect_manuscript.docx"
)
MML2OMML_XSL_PATH = Path(
    r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
)


@dataclass(frozen=True)
class DocumentTokens:
    preset_name: str = "narrative_proposal"
    page_width_in: float = 8.5
    page_height_in: float = 11.0
    margin_in: float = 1.0
    header_footer_distance_in: float = 0.492
    content_width_in: float = 6.5
    base_font: str = "Calibri"
    body_size_pt: float = 11.0
    body_after_pt: float = 8.0
    body_line_spacing: float = 1.333
    title_size_pt: float = 20.0
    title_color: str = "0B2545"
    heading_1_size_pt: float = 16.0
    heading_2_size_pt: float = 13.0
    heading_3_size_pt: float = 12.0
    heading_blue: str = "2E74B5"
    heading_dark_blue: str = "1F4D78"
    table_width_dxa: int = 9360
    table_indent_dxa: int = 120
    table_header_fill: str = "F4F6F9"
    table_cell_top_bottom_dxa: int = 80
    table_cell_start_end_dxa: int = 120


TOKENS = DocumentTokens()


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    rows: tuple[tuple[str, ...], ...] = ()
    alt: str = ""


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def _set_run_font(run, name: str, size_pt: float | None = None) -> None:
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _configure_styles(document: Document, tokens: DocumentTokens) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = tokens.base_font
    normal.font.size = Pt(tokens.body_size_pt)
    normal._element.rPr.rFonts.set(qn("w:ascii"), tokens.base_font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), tokens.base_font)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), tokens.base_font)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(tokens.body_after_pt)
    normal.paragraph_format.line_spacing = tokens.body_line_spacing

    heading_specs = (
        ("Heading 1", tokens.heading_1_size_pt, tokens.heading_blue, 18, 10),
        ("Heading 2", tokens.heading_2_size_pt, tokens.heading_blue, 12, 6),
        ("Heading 3", tokens.heading_3_size_pt, tokens.heading_dark_blue, 8, 4),
    )
    for name, size, color, before, after in heading_specs:
        style = styles[name]
        style.font.name = tokens.base_font
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), tokens.base_font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), tokens.base_font)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), tokens.base_font)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = tokens.base_font
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = _rgb("4B5563")
    caption._element.rPr.rFonts.set(qn("w:ascii"), tokens.base_font)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), tokens.base_font)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), tokens.base_font)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = tokens.base_font
        style.font.size = Pt(tokens.body_size_pt)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    _set_run_font(run, TOKENS.base_font, 8.5)


def _configure_page(document: Document, tokens: DocumentTokens) -> None:
    section = document.sections[0]
    section.page_width = Inches(tokens.page_width_in)
    section.page_height = Inches(tokens.page_height_in)
    section.top_margin = Inches(tokens.margin_in)
    section.right_margin = Inches(tokens.margin_in)
    section.bottom_margin = Inches(tokens.margin_in)
    section.left_margin = Inches(tokens.margin_in)
    section.header_distance = Inches(tokens.header_footer_distance_in)
    section.footer_distance = Inches(tokens.header_footer_distance_in)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Support-aware distributional effect analysis")
    _set_run_font(header_run, tokens.base_font, 8.5)
    header_run.font.color.rgb = _rgb("6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_field(footer)


def _iter_blocks(text: str) -> Iterator[MarkdownBlock]:
    lines = text.splitlines()
    index = 0
    in_comment = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if in_comment:
            if "-->" in stripped:
                in_comment = False
            index += 1
            continue
        if stripped.startswith("<!--"):
            in_comment = "-->" not in stripped
            index += 1
            continue
        if not stripped:
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            yield MarkdownBlock("heading", heading.group(2), len(heading.group(1)))
            index += 1
            continue

        image = re.match(r"^!\[([^]]*)\]\(([^)]+)\)$", stripped)
        if image:
            yield MarkdownBlock("image", image.group(2), alt=image.group(1))
            index += 1
            continue

        if stripped == r"\[":
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != r"\]":
                equation_lines.append(lines[index].strip())
                index += 1
            index += 1
            yield MarkdownBlock("equation", " ".join(equation_lines))
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = []
            for table_line in table_lines:
                cells = tuple(cell.strip() for cell in table_line.strip("|").split("|"))
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            yield MarkdownBlock("table", rows=tuple(rows))
            continue

        list_match = re.match(r"^([-*]|\d+\.)\s+(.+)$", stripped)
        if list_match:
            marker, item = list_match.groups()
            index += 1
            continuation = []
            while index < len(lines):
                candidate = lines[index]
                if not candidate.strip() or re.match(r"^\s*([-*]|\d+\.)\s+", candidate):
                    break
                if candidate.startswith(" "):
                    continuation.append(candidate.strip())
                    index += 1
                else:
                    break
            yield MarkdownBlock("number" if marker.endswith(".") else "bullet", " ".join((item, *continuation)))
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "|", "![", "<!--"))
                or candidate == r"\["
                or re.match(r"^([-*]|\d+\.)\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        yield MarkdownBlock("paragraph", " ".join(paragraph_lines))


INLINE_PATTERN = re.compile(
    r"(\\\(.+?\\\)|\*\*.+?\*\*|`.+?`|\*[^*]+?\*|\[[^]]+\]\([^)]+\))"
)


def _add_inline_runs(paragraph, text: str, size_pt: float | None = None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _set_run_font(run, TOKENS.base_font, size_pt)
        token = match.group(0)
        if token.startswith(r"\("):
            paragraph._p.append(_latex_to_omml(token[2:-2]))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            _set_run_font(run, TOKENS.base_font, size_pt)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, "Consolas", size_pt or 9.5)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            _set_run_font(run, TOKENS.base_font, size_pt)
        else:
            label, url = re.match(r"^\[([^]]+)\]\(([^)]+)\)$", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            _set_run_font(run, TOKENS.base_font, size_pt)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run, TOKENS.base_font, size_pt)


def _set_cell_margins(cell, tokens: DocumentTokens) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", tokens.table_cell_top_bottom_dxa),
        ("bottom", tokens.table_cell_top_bottom_dxa),
        ("start", tokens.table_cell_start_end_dxa),
        ("end", tokens.table_cell_start_end_dxa),
    ):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _table_column_widths(rows: tuple[tuple[str, ...], ...], total_dxa: int) -> list[int]:
    column_count = len(rows[0])
    weights = []
    for column in range(column_count):
        values = [row[column] for row in rows]
        numeric = all(re.fullmatch(r"[-+0-9.%]+", value.replace(" ", "")) for value in values[1:])
        longest = max(len(value) for value in values)
        weights.append(9 if numeric else max(12, min(45, longest)))
    width_sum = sum(weights)
    widths = [round(total_dxa * weight / width_sum) for weight in weights]
    widths[-1] += total_dxa - sum(widths)
    return widths


def _set_table_geometry(table, widths: list[int], tokens: DocumentTokens) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(tokens.table_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(tokens.table_indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell, tokens)


def _add_table(document: Document, rows: tuple[tuple[str, ...], ...], tokens: DocumentTokens) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    font_size = 9.2 if len(rows[0]) <= 3 else 8.2 if len(rows[0]) <= 5 else 7.2
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            _add_inline_runs(paragraph, value, font_size)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                _shade_cell(cell, tokens.table_header_fill)
    header_property = OxmlElement("w:tblHeader")
    header_property.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_property)
    _set_table_geometry(table, _table_column_widths(rows, tokens.table_width_dxa), tokens)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_title(document: Document, title: str, tokens: DocumentTokens) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(title)
    _set_run_font(run, tokens.base_font, tokens.title_size_pt)
    run.bold = True
    run.font.color.rgb = _rgb(tokens.title_color)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    meta_run = meta.add_run("Manuscript draft for supervisor review · 16 July 2026")
    _set_run_font(meta_run, tokens.base_font, 9.5)
    meta_run.font.color.rgb = _rgb("6B7280")


def _latex_to_omml(latex: str):
    transform = etree.XSLT(etree.parse(str(MML2OMML_XSL_PATH)))
    mathml = etree.fromstring(latex_to_mathml(latex).encode("utf-8"))
    omml = transform(mathml).getroot()
    math_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    for nary in omml.xpath(".//m:nary[m:e and not(m:e/*)]", namespaces={"m": math_namespace}):
        parent = nary.getparent()
        next_index = parent.index(nary) + 1
        if next_index < len(parent):
            summand = parent[next_index]
            parent.remove(summand)
            nary.find(f"{{{math_namespace}}}e").append(summand)
    return parse_xml(etree.tostring(omml))


def build_document(source_path: Path = SOURCE_PATH, output_path: Path = OUTPUT_PATH) -> Path:
    document = Document()
    _configure_styles(document, TOKENS)
    _configure_page(document, TOKENS)
    document.core_properties.title = "Support-Aware Distributional Effect Analysis for Small-Sample Thermal-Spray Experiments"
    document.core_properties.subject = "Manuscript draft for supervisor review"

    in_references = False
    for block in _iter_blocks(source_path.read_text(encoding="utf-8")):
        if block.kind == "heading":
            if block.level == 1:
                _add_title(document, block.text, TOKENS)
                continue
            if block.text == "References":
                in_references = True
            style = "Heading 1" if block.level == 2 else "Heading 2"
            paragraph = document.add_paragraph(style=style)
            if block.text == "References":
                paragraph.paragraph_format.page_break_before = True
            _add_inline_runs(paragraph, block.text)
            continue

        if block.kind == "image":
            image_path = source_path.parent / block.text
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            picture = paragraph.add_run().add_picture(str(image_path), width=Inches(6.4))
            picture._inline.docPr.set("descr", block.alt)
            picture._inline.docPr.set("title", block.alt)
            continue

        if block.kind == "table":
            _add_table(document, block.rows, TOKENS)
            continue

        if block.kind == "equation":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_together = True
            paragraph._p.append(_latex_to_omml(block.text))
            continue

        if block.kind in {"bullet", "number"}:
            paragraph = document.add_paragraph(style="List Bullet" if block.kind == "bullet" else "List Number")
            _add_inline_runs(paragraph, block.text)
            continue

        text = block.text
        is_figure_caption = text.startswith("**Figure ")
        is_table_caption = text.startswith("**Table ")
        if is_figure_caption or is_table_caption:
            paragraph = document.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_figure_caption else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = is_table_caption
            _add_inline_runs(paragraph, text, 9)
        else:
            paragraph = document.add_paragraph()
            if in_references:
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                paragraph.paragraph_format.space_after = Pt(5)
                paragraph.paragraph_format.line_spacing = 1.0
                _add_inline_runs(paragraph, text, 9.5)
            else:
                _add_inline_runs(paragraph, text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_document())
